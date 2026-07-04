# src/rag_pipeline.py

import os
from typing import List, Generator, TypedDict
import chromadb
from chromadb.config import Settings
import pypdf
import requests
import docx as python_docx
from bs4 import BeautifulSoup
from sentence_transformers import CrossEncoder

from langchain_core.documents.base import Document
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_chroma import Chroma
from langchain_ollama import ChatOllama
from langchain_core.messages import BaseMessage
from langchain_core.output_parsers import StrOutputParser
from langgraph.graph import StateGraph, END

from src.config import (LLM_MODEL, EMBEDDING_MODEL, DOCUMENTS_PATH,
                        VECTOR_DB_PATH, CHUNK_SIZE, CHUNK_OVERLAP,
                        RERANKER_MODEL, RETRIEVER_TOP_K, RERANKER_TOP_N,
                        MAX_REWRITE_ATTEMPTS)
from src.prompts import CONTEXTUALIZE_Q_PROMPT, QA_PROMPT, GRADE_PROMPT, REWRITE_PROMPT


class AgentState(TypedDict):
    question: str
    chat_history: List[BaseMessage]
    standalone_q: str
    documents: List[Document]
    attempts: int


def _read_pdf(filepath: str) -> List[Document]:
    """Read a single PDF file and return one Document per page."""
    docs = []
    try:
        reader = pypdf.PdfReader(filepath)
        for page_num, page in enumerate(reader.pages):
            text = page.extract_text() or ""
            if text.strip():
                docs.append(Document(
                    page_content=text,
                    metadata={"source": filepath, "page": page_num},
                ))
    except Exception as e:
        print(f"Warning: could not read {os.path.basename(filepath)}: {e}")
    return docs


def _read_txt(filepath: str) -> List[Document]:
    """Read a .txt or .md file as a single Document."""
    try:
        with open(filepath, "r", encoding="utf-8") as f:
            text = f.read()
        if text.strip():
            return [Document(page_content=text, metadata={"source": filepath, "page": 0})]
    except Exception as e:
        print(f"Warning: could not read {os.path.basename(filepath)}: {e}")
    return []


def _read_docx(filepath: str) -> List[Document]:
    """Read a .docx file as a single Document."""
    try:
        doc = python_docx.Document(filepath)
        text = "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
        if text.strip():
            return [Document(page_content=text, metadata={"source": filepath, "page": 0})]
    except Exception as e:
        print(f"Warning: could not read {os.path.basename(filepath)}: {e}")
    return []


def _read_url(url: str) -> List[Document]:
    """Fetch a URL and extract its text content as a single Document."""
    try:
        resp = requests.get(url, timeout=15, headers={"User-Agent": "Mozilla/5.0"})
        resp.raise_for_status()
        soup = BeautifulSoup(resp.text, "html.parser")
        for tag in soup(["script", "style", "nav", "footer", "header"]):
            tag.decompose()
        text = soup.get_text(separator="\n", strip=True)
        if text.strip():
            return [Document(page_content=text, metadata={"source": url, "page": 0})]
    except Exception as e:
        print(f"Warning: could not fetch {url}: {e}")
    return []


def _read_file(filepath: str) -> List[Document]:
    """Dispatch to the correct reader based on file extension."""
    ext = os.path.splitext(filepath)[1].lower()
    if ext == ".pdf":
        return _read_pdf(filepath)
    elif ext in (".txt", ".md"):
        return _read_txt(filepath)
    elif ext == ".docx":
        return _read_docx(filepath)
    print(f"Unsupported file type: {ext}")
    return []


def _split_documents(docs: List[Document], chunk_size: int, chunk_overlap: int) -> List[Document]:
    """Split documents into overlapping chunks on natural text boundaries."""
    separators = ["\n\n", "\n", ". ", " ", ""]
    splits = []

    for doc in docs:
        text = doc.page_content
        chunks = _recursive_split(text, chunk_size, chunk_overlap, separators)
        for i, chunk in enumerate(chunks):
            splits.append(Document(
                page_content=chunk,
                metadata={**doc.metadata, "start_index": i},
            ))
    return splits


def _recursive_split(text: str, chunk_size: int, chunk_overlap: int, separators: List[str]) -> List[str]:
    """Recursively split text on decreasing separators until chunks fit chunk_size."""
    if len(text) <= chunk_size:
        return [text] if text.strip() else []

    sep = ""
    remaining = list(separators)
    while remaining:
        candidate = remaining.pop(0)
        if candidate and candidate in text:
            sep = candidate
            break

    if not sep:
        # Fall back to hard character split
        return _hard_split(text, chunk_size, chunk_overlap)

    parts = text.split(sep)
    chunks: List[str] = []
    current = ""

    for part in parts:
        candidate = (current + sep + part).lstrip(sep) if current else part
        if len(candidate) <= chunk_size:
            current = candidate
        else:
            if current:
                chunks.append(current)
            if len(part) > chunk_size:
                chunks.extend(_recursive_split(part, chunk_size, chunk_overlap, remaining[:]))
                current = ""
            else:
                current = part

    if current:
        chunks.append(current)

    # Apply overlap: prepend tail of previous chunk to next
    if chunk_overlap <= 0 or len(chunks) <= 1:
        return [c for c in chunks if c.strip()]

    overlapped: List[str] = [chunks[0]]
    for i in range(1, len(chunks)):
        prev_tail = chunks[i - 1][-chunk_overlap:]
        merged = prev_tail + sep + chunks[i]
        overlapped.append(merged if len(merged) <= chunk_size * 1.2 else chunks[i])

    return [c for c in overlapped if c.strip()]


def _hard_split(text: str, chunk_size: int, chunk_overlap: int) -> List[str]:
    """Simple sliding-window character split when no separator works."""
    chunks = []
    start = 0
    while start < len(text):
        end = min(start + chunk_size, len(text))
        chunks.append(text[start:end])
        if end == len(text):
            break
        start = end - chunk_overlap
    return chunks


class RAGPipeline:
    def __init__(self):
        """Initialize the RAG pipeline components."""

        print(f"Loading LLM model: {LLM_MODEL}...")
        self.llm = ChatOllama(model=LLM_MODEL, temperature=0, keep_alive="5m")

        print(f"Loading Embedding model: {EMBEDDING_MODEL}...")
        self.embeddings = HuggingFaceEmbeddings(
            model_name=EMBEDDING_MODEL,
            model_kwargs={"device": "cpu", "local_files_only": True},
            encode_kwargs={"normalize_embeddings": True}
        )

        print(f"Loading Reranker model: {RERANKER_MODEL}...")
        self.reranker = CrossEncoder(
            RERANKER_MODEL,
            max_length=512,
            device="cpu",
            local_files_only=True,
        )

        print("Initializing ChromaDB Client...")
        self.client = chromadb.PersistentClient(
            path=VECTOR_DB_PATH,
            settings=Settings(allow_reset=True)
        )

        self.vector_store = None
        self.chain = None  # None = not ready; non-None = ready

    def load_and_split_file(self, filepath: str) -> List[Document]:
        """Load and split a single file (PDF, TXT, MD, DOCX)."""
        docs = _read_file(filepath)
        if not docs:
            return []
        splits = _split_documents(docs, CHUNK_SIZE, CHUNK_OVERLAP)
        print(f"Loaded {len(splits)} chunks from {os.path.basename(filepath)}.")
        return splits

    def load_and_split_url(self, url: str) -> List[Document]:
        """Fetch a URL and split its content into chunks."""
        docs = _read_url(url)
        if not docs:
            return []
        splits = _split_documents(docs, CHUNK_SIZE, CHUNK_OVERLAP)
        print(f"Loaded {len(splits)} chunks from {url}.")
        return splits

    def embed_and_store_documents(self, splits: List[Document]) -> bool:
        """Add document chunks to the vector database (additive, no reset)."""
        if not splits:
            print("No splits to process.")
            return False

        if self.vector_store is None:
            self.vector_store = Chroma(
                embedding_function=self.embeddings,
                client=self.client,
                collection_name="paper_navigator_db"
            )

        # Remove existing chunks for the same sources to avoid duplicates on re-upload
        sources = {doc.metadata.get("source", "") for doc in splits}
        for source in sources:
            if source:
                try:
                    self.vector_store._collection.delete(where={"source": source})
                except Exception:
                    pass

        print(f"Embedding {len(splits)} chunks into ChromaDB...")
        self.vector_store.add_documents(splits)
        print("Chunks added to vector database.")

        self.setup_rag_chain()
        return True

    def remove_document(self, filepath: str) -> None:
        """Remove all chunks belonging to a file from the vector database."""
        if self.vector_store is None:
            return
        self.vector_store._collection.delete(where={"source": filepath})
        print(f"Removed {os.path.basename(filepath)} from vector database.")
        if self.vector_store._collection.count() == 0:
            self.chain = None

    def list_indexed_files(self) -> List[str]:
        """Return sorted list of file paths currently indexed in the vector database."""
        if self.vector_store is None:
            return []
        result = self.vector_store._collection.get(include=["metadatas"])
        paths = {m.get("source", "") for m in result["metadatas"] if m.get("source")}
        return sorted(paths)

    def load_existing_index(self) -> bool:
        """Load existing vector database index if available."""
        try:
            collections = self.client.list_collections()
            if not collections:
                print("No existing collections found in DB.")
                return False
        except Exception:
            return False

        print("Loading existing ChromaDB index...")
        self.vector_store = Chroma(
            embedding_function=self.embeddings,
            client=self.client,
            collection_name="paper_navigator_db"
        )

        self.setup_rag_chain()
        return True

    def setup_rag_chain(self):
        """Set up RAG components using pure langchain_core LCEL."""
        if not self.vector_store:
            raise ValueError("Vector store is not initialized.")

        retriever = self.vector_store.as_retriever(
            search_type="similarity",
            search_kwargs={"k": RETRIEVER_TOP_K}
        )

        def format_docs(docs: List[Document]) -> str:
            return "\n\n".join([doc.page_content for doc in docs])

        def contextualize_question(inputs: dict) -> str:
            """Rephrase question as standalone if chat history exists."""
            if not inputs.get("chat_history"):
                return inputs["input"]
            messages = CONTEXTUALIZE_Q_PROMPT.format_messages(
                input=inputs["input"],
                chat_history=inputs["chat_history"]
            )
            return self.llm.invoke(messages).content

        self._retriever = retriever
        self._format_docs = format_docs
        self._contextualize_fn = contextualize_question

        # Answer generation chain: context + history + question → streamed answer
        self._answer_chain = QA_PROMPT | self.llm | StrOutputParser()

        self.chain = self._answer_chain  # non-None signals pipeline is ready
        self._setup_graph()
        print("RAG Chain initialized successfully.")

    def _retrieve_node(self, state: AgentState) -> AgentState:
        """Graph node: contextualize (first pass only) → retrieve → rerank."""
        q = state["standalone_q"] or self._contextualize_fn(
            {"input": state["question"], "chat_history": state["chat_history"]}
        )
        docs = self._rerank(q, self._retriever.invoke(q))
        return {**state, "standalone_q": q, "documents": docs}

    def _grade_docs_node(self, state: AgentState) -> AgentState:
        """Graph node: LLM grades which docs are relevant (1 call for all docs)."""
        docs = state["documents"]
        if not docs:
            return state
        doc_list = "\n\n".join(f"[{i}] {d.page_content[:400]}" for i, d in enumerate(docs))
        messages = GRADE_PROMPT.format_messages(
            question=state["standalone_q"],
            documents=doc_list,
        )
        result = self.llm.invoke(messages).content.strip().lower()
        relevant = []
        if result != "none":
            for part in result.split(","):
                try:
                    idx = int(part.strip())
                    if 0 <= idx < len(docs):
                        relevant.append(docs[idx])
                except ValueError:
                    pass
        return {**state, "documents": relevant}

    def _rewrite_query_node(self, state: AgentState) -> AgentState:
        """Graph node: rewrite standalone_q to improve retrieval on next attempt."""
        messages = REWRITE_PROMPT.format_messages(question=state["standalone_q"])
        new_q = self.llm.invoke(messages).content.strip()
        print(f"Rewritten query: {new_q}")
        return {**state, "standalone_q": new_q, "attempts": state["attempts"] + 1}

    def _should_retry(self, state: AgentState) -> str:
        """Conditional edge: retry if no relevant docs and attempts not exhausted."""
        if not state["documents"] and state["attempts"] < MAX_REWRITE_ATTEMPTS:
            print(f"No relevant docs — rewriting query (attempt {state['attempts'] + 1})")
            return "rewrite"
        return "generate"

    def _setup_graph(self):
        """Compile the Corrective RAG LangGraph."""
        workflow = StateGraph(AgentState)
        workflow.add_node("retrieve", self._retrieve_node)
        workflow.add_node("grade_docs", self._grade_docs_node)
        workflow.add_node("rewrite_query", self._rewrite_query_node)
        workflow.set_entry_point("retrieve")
        workflow.add_edge("retrieve", "grade_docs")
        workflow.add_conditional_edges(
            "grade_docs",
            self._should_retry,
            {"rewrite": "rewrite_query", "generate": END},
        )
        workflow.add_edge("rewrite_query", "retrieve")
        self.graph = workflow.compile()
        print("Agentic RAG graph compiled.")

    def _rerank(self, query: str, docs: List[Document]) -> List[Document]:
        """Rerank docs with cross-encoder, return top RERANKER_TOP_N."""
        if not docs:
            return docs
        pairs = [(query, doc.page_content) for doc in docs]
        scores = self.reranker.predict(pairs)
        ranked = sorted(zip(scores, docs), key=lambda x: x[0], reverse=True)
        return [doc for _, doc in ranked[:RERANKER_TOP_N]]

    def chat(self, question: str, chat_history: List[BaseMessage]) -> Generator:
        """Stream a response, yielding source docs first then answer chunks."""
        if self.chain is None:
            raise ValueError("RAG chain is not initialized.")

        # Run agentic graph: retrieve → grade → (rewrite → retrieve)* → final docs
        final_state = self.graph.invoke({
            "question": question,
            "chat_history": chat_history,
            "standalone_q": "",
            "documents": [],
            "attempts": 0,
        })
        docs = final_state["documents"]

        # Yield source docs so the caller can display citations
        yield {"context": docs}

        # Stream the answer
        for chunk in self._answer_chain.stream({
            "context": self._format_docs(docs),
            "chat_history": chat_history,
            "input": question,
        }):
            yield {"answer": chunk}

    def clear_workspace(self):
        """Clear the entire workspace including documents and vector database."""
        print("Clearing Vector DB...")
        try:
            self.client.reset()
            self.vector_store = None
            self.chain = None
        except Exception as e:
            print(f"Error resetting DB: {e}")

        if os.path.exists(DOCUMENTS_PATH):
            for f in os.listdir(DOCUMENTS_PATH):
                try:
                    os.remove(os.path.join(DOCUMENTS_PATH, f))
                except Exception as e:
                    print(f"Error deleting file {f}: {e}")
        print("Workspace cleared successfully.")
